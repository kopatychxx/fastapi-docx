from fastapi import FastAPI, Form, Request
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from database import SessionLocal, TemplateFile
from docx import Document
import os
import zipfile
import io
from datetime import datetime, timedelta

app = FastAPI()
templates = Jinja2Templates(directory="templates")

# Получаем шаблоны из базы данных
async def get_templates():
    async with SessionLocal() as session:
        result = await session.execute(select(Template))
        return result.scalars().all()

# Форма для ввода данных
@app.get("/", response_class=HTMLResponse)
async def show_form(request: Request):
    return templates.TemplateResponse("form.html", {"request": request})

# Генерация DOCX
@app.post("/generate-docx/")
async def generate_docx(
    client_name: str = Form(...),
    street_address: str = Form(...),
    city_postal_code: str = Form(...),
    client_email: str = Form(...),
    accident_date: str = Form(...),
    current_date: str = Form(...),
    file_number: str = Form(...),
    client_dob: str = Form(...),
    client_abhc: str = Form(...),
    insurance_company_name_section_b: str = Form(...),
    insurance_company_street_section_b: str = Form(...),
    insurance_company_city_postal_section_b: str = Form(...),
    insurance_company_contact_section_b: str = Form(...),
    claim_section_b: str = Form(...),
    insurance_company_name_section_a: str = Form(...),
    insurance_company_street_section_a: str = Form(...),
    insurance_company_city_postal_section_a: str = Form(...),
    insurance_company_contact_section_a: str = Form(...),
    defendant_name: str = Form(...),
    policy_section_a: str = Form(...),
    police_number: str = Form(...),
    collision_number: str = Form(...),
    accident_place: str = Form(...),
    clinic_name: str = Form(...),
    clinic_address_street: str = Form(...),
    clinic_address_city: str = Form(...),
    client_sin: str = Form(...),
    employer_name: str = Form(...),
    employer_address_street: str = Form(...),
    employer_address_city: str = Form(...),
):
    templates_data = await get_templates()

    # Рассчитываем AccidentDateEdit (AccidentDate - 5 лет)
    accident_date_obj = datetime.strptime(accident_date, "%Y-%m-%d")
    accident_date_edit = (accident_date_obj - timedelta(days=5*365)).strftime("%Y-%m-%d")

    # Замены переменных
    replacements = {
        "{{FileNumber}}": file_number,
        "{{CurrentDate}}": current_date,
        "{{ClientName}}": client_name,
        "{{ClientAddressStreet}}": street_address,
        "{{ClientAddressCity}}": city_postal_code,
        "{{ClientEmail}}": client_email,
        "{{AccidentDate}}": accident_date,
        "{{AccidentDateEdit}}": accident_date_edit,
        "{{ClientDOB}}": client_dob,
        "{{ClientABHC}}": client_abhc,
        "{{InsuranceCompanyNameSectionB}}": insurance_company_name_section_b,
        "{{InsuranceCompanyStreetSectionB}}": insurance_company_street_section_b,
        "{{InsuranceCompanyCityPostalSectionB}}": insurance_company_city_postal_section_b,
        "{{InsuranceCompanyContactSectionB}}": insurance_company_contact_section_b,
        "{{ClaimSectionB}}": claim_section_b,
        "{{InsuranceCompanyNameSectionA}}": insurance_company_name_section_a,
        "{{InsuranceCompanyStreetSectionA}}": insurance_company_street_section_a,
        "{{InsuranceCompanyCityPostalSectionA}}": insurance_company_city_postal_section_a,
        "{{InsuranceCompanyContactSectionA}}": insurance_company_contact_section_a,
        "{{DefendantName}}": defendant_name,
        "{{PolicySectionA}}": policy_section_a,
        "{{PoliceNumber}}": police_number,
        "{{CollisionNumber}}": collision_number,
        "{{AccidentPlace}}": accident_place,
        "{{ClinicName}}": clinic_name,
        "{{ClinicAddressStreet}}": clinic_address_street,
        "{{ClinicAddressCity}}": clinic_address_city,
        "{{ClientSin}}": client_sin,
        "{{EmployerName}}": employer_name,
        "{{EmployerAddressStreet}}": employer_address_street,
        "{{EmployerAddressCity}}": employer_address_city,
    }

    # Виртуальная папка для архива
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for template_data in templates_data:
            # Создаём документ из шаблона в памяти
            doc = Document(io.BytesIO(template_data.content))

            # Заменяем переменные в документе
            for para in doc.paragraphs:
                for key, value in replacements.items():
                    if key in para.text:
                        para.text = para.text.replace(key, value)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, value)

            # Записываем документ в архив в памяти
            doc_filename = f"generated_{template_data.filename}"
            with io.BytesIO() as doc_buffer:
                doc.save(doc_buffer)
                zipf.writestr(doc_filename, doc_buffer.getvalue())

    # Получаем содержимое архива
    zip_buffer.seek(0)

    # Отправляем архив в ответ
    return FileResponse(zip_buffer, media_type="application/zip", filename=f"documents_package_{client_name}.zip")
